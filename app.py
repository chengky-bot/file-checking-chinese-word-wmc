# -*- coding: utf-8 -*-
"""
教材/單元報告審核工具（最終版 + 紅色文字預覽 + 明確修正說明）
由程sir設計
"""

import io
import os
import re
from copy import deepcopy
from collections import defaultdict
from typing import Any, Dict, List, Optional, Set, Tuple
import streamlit as st
from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import RGBColor
from datetime import datetime
import fitz  # pymupdf

# ---------------------------------------------------------------------------
# 內建規則
# ---------------------------------------------------------------------------
BUILTIN_REPLACEMENTS: Dict[str, str] = {
    "遺規": "違規",
    "冲動": "衝動",
    "散慢": "散漫",
    "違尿": "遺尿",
    "暴燥": "暴躁",
    "暴騷": "暴躁",
    "書藉": "書籍",
    "加許": "嘉許",
    "座標": "坐標",
    "裡": "裏",
    "溫度": "温度",
    "部份": "部分",
    "盡量": "儘量",
    "掃瞄二維碼": "掃描二維碼",
    "越來越多": "愈來愈多",
    "找贖": "找續",
    "計畫": "計劃",
}

WORD_JOINER = "\u2060"
WORD_GLUE_INNER = WORD_JOINER * 2
_SPLIT_GAP_RE = r"(?:\s*[/／]?\s*)"
_SENTENCE_END_CHARS = set("。！？!?…」』\"'）)】］＞>、，,；;：:")
_JILU_VERB_HINTS = ("了", "著", "過", "在", "中", "時", "一下", "起來", "成", "到", "給", "好", "完", "下", "出")
DEFAULT_INDIVISIBLE = "廣場\n落淚\n靈魂"
IMAGE_PLACEHOLDER = "\ufdd0"

# ---------------------------------------------------------------------------
# 輔助函數（保持不變）
# ---------------------------------------------------------------------------
def _parse_custom_rules(text: str) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for line in text.splitlines():
        line = line.strip()
        if not line or ":" not in line:
            continue
        old, new = line.split(":", 1)
        old, new = old.strip(), new.strip()
        if old:
            out[old] = new
    return out

def _parse_lines(text: str) -> List[str]:
    return [ln.strip() for ln in text.splitlines() if ln.strip()]

def _merge_intervals(intervals: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
    if not intervals:
        return []
    intervals = sorted(intervals)
    merged = [intervals[0]]
    for a, b in intervals[1:]:
        la, lb = merged[-1]
        if a <= lb:
            merged[-1] = (la, max(lb, b))
        else:
            merged.append((a, b))
    return merged

def _intervals_to_set(intervals: List[Tuple[int, int]]) -> Set[int]:
    s: Set[int] = set()
    for a, b in intervals:
        s.update(range(a, b))
    return s

def _replace_non_overlapping(
    text: str,
    pattern_to_repl: List[Tuple[str, str]],
    on_each: Optional[Callable[[int, int, str], None]] = None,
) -> str:
    if not text:
        return text
    result: List[str] = []
    i = 0
    n = len(text)
    while i < n:
        best_j: Optional[int] = None
        best_old = ""
        best_new = ""
        for old, new in pattern_to_repl:
            j = text.find(old, i)
            if j == -1:
                continue
            if best_j is None or j < best_j or (j == best_j and len(old) > len(best_old)):
                best_j = j
                best_old = old
                best_new = new
        if best_j is None:
            result.append(text[i:])
            break
        result.append(text[i:best_j])
        start = len("".join(result))
        result.append(best_new)
        end = start + len(best_new)
        if on_each:
            on_each(start, end, f"{best_old}→{best_new}")
        i = best_j + len(best_old)
    return "".join(result)

def _apply_sorted_replacements(
    text: str,
    replacements: Dict[str, str],
    stats: defaultdict,
    stat_prefix: str,
) -> Tuple[str, List[Tuple[int, int]]]:
    if not replacements:
        return text, []
    pairs = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)
    pattern_to_repl = [(a, b) for a, b in pairs if a]
    red_intervals: List[Tuple[int, int]] = []
    def on_rep(start: int, end: int, label: str) -> None:
        red_intervals.append((start, end))
        stats[f"{stat_prefix}:{label}"] += 1
    new_text = _replace_non_overlapping(text, pattern_to_repl, on_each=on_rep)
    return new_text, _merge_intervals(red_intervals)

def _apply_single_char_replace(
    text: str,
    old_ch: str,
    new_ch: str,
    stats: defaultdict,
    stat_key: str,
) -> Tuple[str, List[Tuple[int, int]]]:
    if old_ch not in text:
        return text, []
    parts: List[str] = []
    red: List[Tuple[int, int]] = []
    pos = 0
    for i, ch in enumerate(text):
        if ch == old_ch:
            parts.append(text[pos:i])
            base = len("".join(parts))
            parts.append(new_ch)
            red.append((base, base + len(new_ch)))
            stats[stat_key] += 1
            pos = i + 1
    parts.append(text[pos:])
    return "".join(parts), _merge_intervals(red)

def _apply_zhuo_to_zhe(text: str, stats: defaultdict) -> Tuple[str, List[Tuple[int, int]]]:
    red: List[Tuple[int, int]] = []
    if "著" not in text:
        return text, red
    out: List[str] = []
    i = 0
    n = len(text)
    while i < n:
        if text[i] != "著":
            out.append(text[i])
            i += 1
            continue
        if text.startswith("著名", i):
            out.append("著名")
            i += 2
            continue
        if text.startswith("著作", i):
            out.append("著作")
            i += 2
            continue
        start = len("".join(out))
        out.append("着")
        red.append((start, start + 1))
        stats["著→着"] += 1
        i += 1
    return "".join(out), _merge_intervals(red)

def _should_be_jilu_verb(text: str, idx: int) -> Optional[bool]:
    if idx + 2 > len(text):
        return None
    c1 = text[idx + 2] if idx + 2 < len(text) else ""
    if not c1:
        return None
    if c1 in ("了", "著", "過", "在", "中", "時", "成", "到", "給", "好", "完", "下", "出"):
        return True
    if c1 in ("的", "和", "與", "及", "等", "、", "，"):
        return False
    rest = text[idx + 2 : idx + 6]
    for hint in _JILU_VERB_HINTS:
        if len(hint) >= 2 and rest.startswith(hint):
            return True
    return None

def _apply_jilu_jilu(text: str, stats: defaultdict) -> Tuple[str, List[Tuple[int, int]]]:
    red: List[Tuple[int, int]] = []
    if "記錄" not in text and "紀錄" not in text:
        return text, red
    out: List[str] = []
    i = 0
    n = len(text)
    while i < n:
        if i + 1 < n and text[i : i + 2] in ("記錄", "紀錄"):
            pair = text[i : i + 2]
            decision = _should_be_jilu_verb(text, i)
            if decision is None:
                out.append(pair)
                stats["記錄／紀錄待確認"] += 1
                i += 2
                continue
            want = "記錄" if decision else "紀錄"
            if pair == want:
                out.append(pair)
                i += 2
                continue
            start = len("".join(out))
            out.append(want)
            red.append((start, start + 2))
            stats[f"記錄／紀錄→{want}"] += 1
            i += 2
            continue
        out.append(text[i])
        i += 1
    return "".join(out), _merge_intervals(red)

def _apply_wei_to_wei_transition(text: str, stats: defaultdict) -> Tuple[str, List[Tuple[int, int]]]:
    red: List[Tuple[int, int]] = []
    pattern = re.compile(r"，唯(?![一])")
    out: List[str] = []
    last = 0
    for m in pattern.finditer(text):
        out.append(text[last : m.start()])
        start = len("".join(out))
        chunk = "，惟"
        out.append(chunk)
        red.append((start + 1, start + len(chunk)))
        stats["，唯→，惟"] += 1
        last = m.end()
    out.append(text[last:])
    return "".join(out), _merge_intervals(red)

def _should_add_period_at_end(text: str) -> bool:
    s = text.rstrip()
    if len(s) < 8:
        return False
    last = s[-1]
    if last in _SENTENCE_END_CHARS or last.isspace():
        return False
    if last.isalnum() or "\u4e00" <= last <= "\u9fff":
        return True
    return False

def _merge_red(*interval_lists: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
    flat: List[Tuple[int, int]] = []
    for lst in interval_lists:
        flat.extend(lst)
    return _merge_intervals(flat)

def _flex_pattern_for_word(w: str) -> str:
    if len(w) < 2:
        return re.escape(w)
    parts: List[str] = [re.escape(w[0])]
    for ch in w[1:]:
        parts.append(_SPLIT_GAP_RE)
        parts.append(re.escape(ch))
    return "".join(parts)

def _normalize_indivisible_splits(text: str, words: List[str], stats: defaultdict) -> str:
    words_sorted = sorted(set(w for w in words if len(w) >= 2), key=len, reverse=True)
    if not words_sorted:
        return text
    out: List[str] = []
    i = 0
    n = len(text)
    while i < n:
        matched = False
        for w in words_sorted:
            pat = _flex_pattern_for_word(w)
            m = re.match(pat, text[i:])
            if not m:
                continue
            raw = m.group(0)
            out.append(w)
            if raw != w:
                stats["不可分割_合併人工斷字"] += 1
            i += len(raw)
            matched = True
            break
        if not matched:
            out.append(text[i])
            i += 1
    return "".join(out)

def _apply_proper_nouns(
    text: str,
    names: List[str],
    stats: defaultdict,
) -> Tuple[str, List[Tuple[int, int]]]:
    names_sorted = sorted(set(names), key=len, reverse=True)
    underline: List[Tuple[int, int]] = []
    occupied: List[Tuple[int, int]] = []
    def overlaps(a: int, b: int) -> bool:
        for x, y in occupied:
            if not (b <= x or a >= y):
                return True
        return False
    for name in names_sorted:
        if not name:
            continue
        start = 0
        while True:
            j = text.find(name, start)
            if j == -1:
                break
            end = j + len(name)
            if not overlaps(j, end):
                underline.append((j, end))
                occupied.append((j, end))
                stats["專名號底線"] += 1
            start = j + 1
    return text, _merge_intervals(underline)

def _apply_word_joiners(
    text: str,
    words: List[str],
    stats: defaultdict,
) -> Tuple[str, List[Tuple[int, int]]]:
    words_sorted = sorted(set(w for w in words if len(w) >= 2), key=len, reverse=True)
    if not words_sorted:
        return text, []
    occupied: List[Tuple[int, int]] = []
    matches: List[Tuple[int, int, str]] = []
    i = 0
    n = len(text)
    while i < n:
        found: Optional[str] = None
        for w in words_sorted:
            if i + len(w) <= n and text[i : i + len(w)] == w and not overlaps(i, i + len(w)):
                found = w
                break
        if found:
            matches.append((i, i + len(found), found))
            occupied.append((i, i + len(found)))
            i += len(found)
        else:
            i += 1
    if not matches:
        return text, []
    red: List[Tuple[int, int]] = []
    out: List[str] = []
    last = 0
    for start, end, w in matches:
        out.append(text[last:start])
        pos = len("".join(out))
        core = WORD_GLUE_INNER.join(list(w))
        if start > 0:
            piece = WORD_JOINER + core
        else:
            piece = core
        out.append(piece)
        red.append((pos, pos + len(piece)))
        stats["排位調整_WORD_JOINER"] += 1
        last = end
    out.append(text[last:])
    return "".join(out), _merge_intervals(red)

def _map_span_before_to_after(
    before: str,
    after: str,
    bs: int,
    be: int,
) -> Optional[Tuple[int, int]]:
    if bs >= be:
        return (0, 0)
    ib, ia = 0, 0
    nb, na = len(before), len(after)
    a0: Optional[int] = None
    while ib < nb and ia < na:
        while ia < na and after[ia] == WORD_JOINER:
            ia += 1
        if ia >= na:
            break
        if ib == bs:
            a0 = ia
        if ib == be - 1:
            return (a0 if a0 is not None else ia, ia + 1)
        ib += 1
        ia += 1
    return None

def _map_underline_after_joiners(
    text_before: str,
    text_after: str,
    underline_before: List[Tuple[int, int]],
) -> List[Tuple[int, int]]:
    if not underline_before:
        return []
    spans: List[Tuple[int, int]] = []
    for bs, be in _merge_intervals(underline_before):
        m = _map_span_before_to_after(text_before, text_after, bs, be)
        if m:
            spans.append(m)
    return _merge_intervals(spans)

def _paragraph_plain(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)

def _run_has_drawing(run) -> bool:
    xml = run._element
    for el in xml.iter():
        tag = el.tag
        if tag.endswith("drawing") or tag.endswith("pict") or tag.endswith("binaryData"):
            return True
    return False

def _paragraph_text_and_drawings(paragraph) -> Tuple[str, List[Any]]:
    drawings: List[Any] = []
    parts: List[str] = []
    for run in paragraph.runs:
        if _run_has_drawing(run):
            parts.append(IMAGE_PLACEHOLDER)
            drawings.append(deepcopy(run._element))
        else:
            parts.append(run.text or "")
    return "".join(parts), drawings

def _clear_paragraph_content_keep_ppr(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)

def _append_formatted_runs(
    paragraph,
    text: str,
    red_indices: Set[int],
    underline_indices: Set[int],
) -> None:
    if not text:
        return
    n = len(text)
    i = 0
    while i < n:
        is_red = i in red_indices
        is_ul = i in underline_indices
        j = i + 1
        while j < n:
            jr = j in red_indices
            ju = j in underline_indices
            if jr != is_red or ju != is_ul:
                break
            j += 1
        run = paragraph.add_run(text[i:j])
        if is_red:
            run.font.color.rgb = RGBColor(255, 0, 0)
        if is_ul:
            run.font.underline = WD_UNDERLINE.SINGLE
        i = j

def _write_formatted_runs_full(
    paragraph,
    text: str,
    red_indices: Set[int],
    underline_indices: Set[int],
) -> None:
    _clear_paragraph_content_keep_ppr(paragraph)
    _append_formatted_runs(paragraph, text, red_indices, underline_indices)

def _rebuild_paragraph_with_image_placeholders(
    paragraph,
    new_full: str,
    drawing_elements: List[Any],
    red_indices: Set[int],
    underline_indices: Set[int],
) -> bool:
    ph = IMAGE_PLACEHOLDER
    n_ph = new_full.count(ph)
    if n_ph != len(drawing_elements):
        return False
    parts = new_full.split(ph)
    _clear_paragraph_content_keep_ppr(paragraph)
    pos = 0
    ph_len = len(ph)
    for i, part in enumerate(parts):
        seg_start = pos
        seg_end = pos + len(part)
        rs = {x - seg_start for x in red_indices if seg_start <= x < seg_end}
        us = {x - seg_start for x in underline_indices if seg_start <= x < seg_end}
        _append_formatted_runs(paragraph, part, rs, us)
        pos = seg_end
        if i < len(drawing_elements):
            paragraph._p.append(drawing_elements[i])
            pos += ph_len
    return True

def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nt in cell.tables:
                    yield from walk_table(nt)
    for table in doc.tables:
        yield from walk_table(table)
    for section in doc.sections:
        for part in (section.header, section.footer):
            try:
                for p in part.paragraphs:
                    yield p
                for table in part.tables:
                    yield from walk_table(table)
            except Exception:
                pass

def process_paragraph_plain_text(
    raw: str,
    custom_rules: Dict[str, str],
    proper_names: List[str],
    indivisible: List[str],
    stats: defaultdict,
    add_period: bool,
) -> Tuple[str, Set[int], Set[int], List[Dict]]:
    t = raw
    findings: List[Dict] = []

    t, r1 = _apply_sorted_replacements(t, BUILTIN_REPLACEMENTS, stats, "內建")
    t, r2 = _apply_sorted_replacements(t, custom_rules, stats, "自訂")
    t, r3 = _apply_single_char_replace(t, "佈", "布", stats, "佈→布")
    t, r4 = _apply_zhuo_to_zhe(t, stats)
    t, r5 = _apply_jilu_jilu(t, stats)
    t, r6 = _apply_wei_to_wei_transition(t, stats)
    qingjing: Dict[str, str] = {"情景": "情境"}
    t, r7 = _apply_sorted_replacements(t, qingjing, stats, "情景")

    red = _merge_red(r1, r2, r3, r4, r5, r6, r7)

    if add_period and _should_add_period_at_end(t):
        t = t.rstrip() + "。"
        pos = len(t) - 1
        red.append((pos, pos + 1))
        stats["句尾標點補上"] += 1
        findings.append({
            "rule": "句尾標點補上",
            "before": raw,
            "after": t,
            "start": pos,
            "end": pos + 1,
        })

    t = _normalize_indivisible_splits(t, indivisible, stats)
    t_before_join = t
    t, r_join = _apply_word_joiners(t, indivisible, stats)
    red = _merge_red(red, r_join)

    t, underline_before = _apply_proper_nouns(t, proper_names, stats)
    underline_after = _map_underline_after_joiners(t_before_join, t, underline_before)

    red_set = _intervals_to_set(red)
    under_set = _intervals_to_set(underline_after)

    return t, red_set, under_set, findings

def process_document(
    doc: Document,
    custom_rules: Dict[str, str],
    proper_names: List[str],
    indivisible: List[str],
    add_period: bool,
) -> Tuple[defaultdict, List[Dict]]:
    stats: defaultdict = defaultdict(int)
    all_findings: List[Dict] = []
    para_index = 0

    for p in iter_all_paragraphs(doc):
        para_index += 1
        location = f"第 {para_index} 段"
        text_ph, drs = _paragraph_text_and_drawings(p)

        new_text, red, under, para_findings = process_paragraph_plain_text(
            text_ph, custom_rules, proper_names, indivisible, stats, add_period
        )

        for f in para_findings:
            f["location"] = location
            f["snippet"] = text_ph[:80] + "..." if len(text_ph) > 80 else text_ph
            all_findings.append(f)

        if not drs:
            _write_formatted_runs_full(p, new_text, red, under)
        else:
            ok = _rebuild_paragraph_with_image_placeholders(p, new_text, drs, red, under)
            if not ok:
                stats["含圖段落占位異常_未改寫"] += 1

    return stats, all_findings

def _total_changes(stats: defaultdict) -> int:
    return int(sum(stats.values()))

# ---------------------------------------------------------------------------
# 主程式
# ---------------------------------------------------------------------------
def main() -> None:
    st.set_page_config(page_title="教材/單元報告審核工具", layout="wide")
    st.title("📚 教材／單元報告審核工具")
    st.markdown("**由程sir設計**")

    # Session State
    if "processed_bytes" not in st.session_state:
        st.session_state.processed_bytes = None
    if "last_stats" not in st.session_state:
        st.session_state.last_stats = None
    if "last_findings" not in st.session_state:
        st.session_state.last_findings = []
    if "report_text" not in st.session_state:
        st.session_state.report_text = None
    if "preview_html" not in st.session_state:
        st.session_state.preview_html = None
    if "download_filename" not in st.session_state:
        st.session_state.download_filename = "document_fixed.docx"

    col_main, col_side = st.columns([3, 1])

    with col_side:
        st.subheader("⚙️ 規則與清單")
        tab1, tab2, tab3 = st.tabs(["📝 自訂規則", "🏷️ 專有名詞", "🔗 不可分割詞"])
        with tab1:
            custom_rules_text = st.text_area(
                "自訂額外規則（每行：舊詞:新詞）",
                height=180,
                placeholder="遺規:違規\n暴燥:暴躁",
            )
        with tab2:
            proper_text = st.text_area(
                "自訂專有名詞清單（每行一個）",
                height=180,
                placeholder="香港大學\n教學樓\n李老師",
            )
        with tab3:
            indiv_text = st.text_area(
                "不可分割詞彙清單（每行一個）",
                value=DEFAULT_INDIVISIBLE,
                height=180,
                placeholder="廣場\n落淚\n靈魂",
            )

        st.divider()
        add_period = st.checkbox(
            "自動補句尾標點",
            value=False,
            help="有些文件（如表格、短句）不想自動補句號，請取消勾選",
        )

    with col_main:
        input_mode = st.radio(
            "選擇輸入方式",
            options=["📤 上傳檔案（DOCX / PDF）", "📋 直接貼上文字"],
            horizontal=True,
        )

        if input_mode == "📤 上傳檔案（DOCX / PDF）":
            up = st.file_uploader("上傳檔案", type=["docx", "pdf"])
            pasted_text = None
        else:
            pasted_text = st.text_area(
                "請在此貼上要檢查的文字",
                height=250,
                placeholder="直接貼上文字後，按下方按鈕",
            )
            up = None

        run_btn = st.button("🚀 開始審核並套用修正", type="primary", use_container_width=True)

    if run_btn:
        with st.spinner("🔍 正在審核，請稍等..."):
            custom = _parse_custom_rules(custom_rules_text)
            names = _parse_lines(proper_text)
            indiv = _parse_lines(indiv_text)

            if input_mode == "📤 上傳檔案（DOCX / PDF）" and up is not None:
                file_bytes = up.getvalue()
                file_name = up.name.lower()

                if file_name.endswith(".pdf"):
                    pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
                    full_text = "\n\n".join(page.get_text("text") for page in pdf_doc)
                    doc = Document()
                    for line in full_text.split("\n"):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                    input_name = up.name
                else:
                    doc = Document(io.BytesIO(file_bytes))
                    input_name = up.name

            elif input_mode == "📋 直接貼上文字" and pasted_text:
                doc = Document()
                for line in pasted_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                input_name = "貼上文字.docx"

            else:
                st.warning("請上傳檔案或貼上文字")
                st.stop()

            stats, findings = process_document(doc, custom, names, indiv, add_period)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.session_state.processed_bytes = buf.getvalue()
            st.session_state.last_stats = dict(stats)
            st.session_state.last_findings = findings
            base = os.path.splitext(input_name)[0]
            st.session_state.download_filename = f"{base}_fixed.docx"

            # ── 產生紅色高亮預覽 ──
            preview_html = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    preview_html += f"<p>{para.text}</p>"
            st.session_state.preview_html = preview_html

            # 產生報告
            report_lines = [
                f"修正報告 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                f"輸入方式：{input_mode}",
                f"檔案／文字：{input_name}",
                f"總修正項目：{_total_changes(defaultdict(int, stats))} 項",
                f"自動補句尾標點：{'已開啟' if add_period else '已關閉'}",
                "=" * 60,
                "",
            ]
            for f in findings:
                start = f.get("start", 0) + 1
                end = f.get("end", start)
                report_lines.append(f"{f.get('location', '')}  第 {start}～{end} 字  {f.get('rule', '')}")
                report_lines.append(f"原本：{f.get('before', '')[:120]}...")
                report_lines.append(f"修正後：{f.get('after', '')[:120]}...")
                report_lines.append("-" * 50)
            st.session_state.report_text = "\n".join(report_lines)

    if st.session_state.processed_bytes is not None:
        st.success(f"✅ 審核完成！總共修正 **{_total_changes(defaultdict(int, st.session_state.last_stats))}** 項")

        col1, col2 = st.columns(2)

        with col1:
            with st.expander("📊 修正統計", expanded=True):
                stats = st.session_state.last_stats or {}
                findings = st.session_state.get("last_findings", [])

                builtin_lines = [f"- {k}: {v}" for k, v in sorted(stats.items()) if k.startswith("內建:")]
                if builtin_lines:
                    st.markdown("**內建規則**")
                    st.markdown("\n".join(builtin_lines))

                custom_lines = [f"- {k}: {v}" for k, v in sorted(stats.items()) if k.startswith("自訂:")]
                if custom_lines:
                    st.markdown("**自訂規則**")
                    st.markdown("\n".join(custom_lines))

                st.markdown("**其他規則**")
                rule_locations = defaultdict(list)
                for f in findings:
                    rule = f.get("rule", "未知規則")
                    loc = f.get("location", "未知段落")
                    rule_locations[rule].append(loc)

                for rule, locs in sorted(rule_locations.items()):
                    count = len(locs)
                    unique_locs = list(dict.fromkeys(locs))
                    st.markdown(f"**• {rule}**　`{count} 次`")
                    loc_str = "、".join(unique_locs[:8])
                    st.caption(f"出現於：{loc_str}")
                    if len(unique_locs) > 8:
                        st.caption(f"... 還有 {len(unique_locs)-8} 段")

        with col2:
            with st.expander("📍 字級問題位置", expanded=True):
                findings = st.session_state.get("last_findings", [])
                if findings:
                    for item in findings:
                        start = item.get("start", 0) + 1
                        end = item.get("end", start)
                        st.markdown(f"**{item.get('location')}　第 {start}～{end} 字** **{item.get('rule')}**")
                        st.caption(f"原本：{item.get('before', '')[:80]}...")
                        st.caption(f"修正後：{item.get('after', '')[:80]}...")
                else:
                    st.info("本次沒有發現需要修正的問題")

        # ── 已修正文字預覽（紅色高亮 + 修正說明） ──
        with st.expander("📖 已修正文字預覽（紅色高亮 + 修正說明）", expanded=True):
            if st.session_state.get("preview_html"):
                # 紅色高亮預覽
                st.markdown(
                    f"""
                    <div style="background-color: #f9f9f9; padding: 20px; border-radius: 8px; border: 1px solid #ddd; font-size: 16px; line-height: 1.7;">
                        {st.session_state.preview_html}
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                st.caption("💡 紅色文字 = 已自動修正的錯別字／用字")

                # 修正說明列表
                findings = st.session_state.get("last_findings", [])
                if findings:
                    st.markdown("### 📋 本次修正的詳細位置")
                    for item in findings:
                        start = item.get("start", 0) + 1
                        end = item.get("end", start)
                        st.markdown(f"**{item.get('location')}　第 {start}～{end} 字** **{item.get('rule')}**")
                        st.caption(f"原本 → 修正後：{item.get('before', '')[:80]}... → {item.get('after', '')[:80]}...")
                else:
                    st.info("本次沒有發現需要修正的問題")
            else:
                st.info("尚未產生預覽文字")

        st.divider()

        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button(
                label="📄 下載已修改的 DOCX",
                data=st.session_state.processed_bytes,
                file_name=st.session_state.download_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_dl2:
            if st.session_state.get("report_text"):
                st.download_button(
                    label="📝 下載修正報告（純文字 .txt）",
                    data=st.session_state.report_text,
                    file_name="修正報告.txt",
                    mime="text/plain",
                    use_container_width=True,
                )

    if st.button("🔄 重置所有設定", use_container_width=True):
        st.session_state.clear()
        st.rerun()

if __name__ == "__main__":
    main()
